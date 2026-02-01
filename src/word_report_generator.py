"""
Word document report generator for easy editing.

Creates an editable .docx version of the technical report.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from pathlib import Path
from datetime import datetime


class WordReportGenerator:
    """Generates an editable Word document version of the technical report."""

    def __init__(self, results_path: str = "outputs/evaluation_results.json"):
        """Initialize with results file."""
        self.results_path = Path(results_path)
        with open(self.results_path, 'r') as f:
            self.results = json.load(f)

        self.figures_dir = Path("outputs/figures")
        self.doc = Document()

        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_horizontal_line(self):
        """Add a horizontal line separator."""
        p = self.doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_before = Pt(6)
        p_format.space_after = Pt(6)

        # Add bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_styled_heading(self, text, level=1, color=None):
        """Add a styled heading."""
        heading = self.doc.add_heading(text, level=level)
        if color:
            for run in heading.runs:
                run.font.color.rgb = color
        return heading

    def generate(self, output_path: str = "outputs/technical_report.docx"):
        """Generate the complete Word report."""

        # Title Page
        title = self.doc.add_heading('TECHNICAL REPORT', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(24)
            run.font.bold = True

        subtitle = self.doc.add_paragraph('Evaluation of Generative Agent Human Simulations')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(14)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

        date_p = self.doc.add_paragraph(f'Evaluation Date: {datetime.now().strftime("%B %d, %Y")}')
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in date_p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(128, 128, 128)

        self._add_horizontal_line()

        # Executive Summary
        self._add_styled_heading('EXECUTIVE SUMMARY', level=1, color=RGBColor(0, 0, 0))

        meta = self.results['metadata']
        agg = self.results['aggregate_scores']

        summary_text = f"""This report evaluates the performance of LLM-based human simulations created for a Beauty & Wellbeing company's market research. The evaluation assessed {meta['total_pairs']} response pairs from {meta['num_persons']} individuals across {meta['num_categories']} question categories using a multi-dimensional framework.

KEY FINDINGS:
• Semantic Alignment: {agg['semantic']['mean_cosine_similarity']:.3f} cosine similarity (0-1 scale)
• Stylistic Consistency: {agg['stylistic']['mean_style_score']:.3f} overall style match (0-1 scale)"""

        if agg['llm_judge']['count'] > 0:
            summary_text += f"\n• LLM Judge Assessment: {agg['llm_judge']['mean_overall_score']:.2f}/10 overall score"
            summary_text += f"\n  - Semantic Match: {agg['llm_judge']['mean_semantic_match']:.2f}/10"
            summary_text += f"\n  - Style Match: {agg['llm_judge']['mean_style_match']:.2f}/10"
            summary_text += f"\n  - Personality Match: {agg['llm_judge']['mean_personality_match']:.2f}/10"
            summary_text += f"\n  - Naturalness: {agg['llm_judge']['mean_naturalness']:.2f}/10"

        summary_text += f"""
• CRITICAL ISSUE: AI responses are {agg['stylistic']['mean_length_ratio']:.0%} longer than human responses
• Only {agg['stylistic']['pct_with_imperfections']:.0f}% of AI responses exhibit human-like imperfections

BENCHMARK CONTEXT:
The referenced paper "Generative Agent Simulations of 1,000 People" achieved 85% accuracy compared to human test-retest reliability (equivalent to 8.5/10). This evaluation uses this as a gold standard."""

        p = self.doc.add_paragraph(summary_text)
        p.paragraph_format.line_spacing = 1.15

        # Page break
        self.doc.add_page_break()

        # Methodology
        self._add_styled_heading('EVALUATION METHODOLOGY', level=1, color=RGBColor(0, 0, 0))

        methodology_text = """This evaluation employs a comprehensive multi-dimensional framework:

1. SEMANTIC SIMILARITY
   Measures content alignment using state-of-the-art sentence embeddings (all-MiniLM-L6-v2).
   Captures whether the AI conveys the same meaning as the human response.

2. STYLISTIC ANALYSIS
   Analyzes linguistic features including response length, readability scores, formality levels,
   and presence of human-like imperfections (typos, casual language).

3. LLM-AS-JUDGE EVALUATION
   Uses Claude Opus 4 to provide holistic expert assessment across multiple dimensions:
   semantic match, style match, personality consistency, and naturalness."""

        self.doc.add_paragraph(methodology_text)

        # Visualizations
        self._add_styled_heading('EVALUATION RESULTS', level=1, color=RGBColor(0, 0, 0))

        # Add overview chart
        if (self.figures_dir / "overview_scores.png").exists():
            self.doc.add_paragraph('Multi-Dimensional Performance Overview', style='Heading 2')
            self.doc.add_picture(str(self.figures_dir / "overview_scores.png"), width=Inches(6))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add dimension comparison
        if (self.figures_dir / "dimension_comparison.png").exists():
            self.doc.add_paragraph('Performance Across Evaluation Dimensions', style='Heading 2')
            self.doc.add_picture(str(self.figures_dir / "dimension_comparison.png"), width=Inches(6))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Per-person analysis
        if (self.figures_dir / "person_performance.png").exists():
            self.doc.add_paragraph('Per-Person Simulation Performance', style='Heading 2')
            self.doc.add_picture(str(self.figures_dir / "person_performance.png"), width=Inches(6))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_page_break()

        # Category performance
        if (self.figures_dir / "category_performance.png").exists():
            self.doc.add_paragraph('Performance by Question Category', style='Heading 2')
            self.doc.add_picture(str(self.figures_dir / "category_performance.png"), width=Inches(6))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Length analysis
        if (self.figures_dir / "length_analysis.png").exists():
            self.doc.add_paragraph('Response Length Analysis', style='Heading 2')
            self.doc.add_picture(str(self.figures_dir / "length_analysis.png"), width=Inches(6))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Key Insights
        self._add_styled_heading('KEY INSIGHTS & WEAKNESSES', level=1, color=RGBColor(0, 0, 0))

        person_analysis = self.results['per_person_analysis']
        best_person = max(person_analysis.items(), key=lambda x: x[1]['mean_semantic_similarity'])

        self.doc.add_paragraph('STRENGTHS', style='Heading 2')
        strengths = [
            f"Semantic content generally well-captured (mean: {agg['semantic']['mean_cosine_similarity']:.3f})",
            f"Best simulation: {best_person[0]} (semantic: {best_person[1]['mean_semantic_similarity']:.3f})"
        ]
        for strength in strengths:
            p = self.doc.add_paragraph(strength, style='List Bullet')

        self.doc.add_paragraph('CRITICAL WEAKNESSES', style='Heading 2')
        weaknesses = [
            f"LENGTH MISMATCH: AI responses average {agg['stylistic']['mean_length_ratio']:.2f}x longer than humans",
            f"LACK OF NATURALNESS: Low human-like imperfection rate ({agg['stylistic']['pct_with_imperfections']:.0f}%)",
            "OVER-POLISHED: AI responses lack casual language, typos, and conversational quirks"
        ]

        if agg['llm_judge']['count'] > 0 and agg['llm_judge']['common_weaknesses']:
            for weakness in agg['llm_judge']['common_weaknesses'][:3]:
                weaknesses.append(f"COMMON PATTERN: {weakness['weakness']}")

        for weakness in weaknesses:
            p = self.doc.add_paragraph(weakness, style='List Bullet')

        # Worst match example
        self.doc.add_paragraph('WORST MATCH EXAMPLE', style='Heading 2')
        weakest = self.results['weakest_matches'][0]

        example_text = f"""ID: {weakest['pair_id']} | Category: {weakest['category']} | Composite Score: {weakest['composite_score']:.3f}

Question: {weakest['question']}

Human Response: "{weakest['human_answer']}"

AI Response: "{weakest['ai_answer']}"

Issues Identified: {', '.join(weakest['issues'][:3]) if weakest['issues'] else 'See detailed analysis'}"""

        p = self.doc.add_paragraph(example_text)
        p.paragraph_format.line_spacing = 1.3

        # Recommendations
        self._add_styled_heading('RECOMMENDATIONS FOR IMPROVEMENT', level=1, color=RGBColor(0, 0, 0))

        recommendations = [
            "RESPONSE LENGTH CALIBRATION: Train simulations to match human verbosity patterns (target: 1.0-1.2x vs current {:.1f}x)".format(agg['stylistic']['mean_length_ratio']),
            "INJECT NATURALNESS: Add human-like imperfections (typos, casual language, filler words)",
            "PERSONALITY FINE-TUNING: Improve individual voice capture through more targeted prompting",
            "CATEGORY-SPECIFIC TUNING: Address performance gaps in underperforming categories (especially 'influences')",
            "ITERATIVE REFINEMENT: Use evaluation feedback to retrain simulation prompts"
        ]

        for i, rec in enumerate(recommendations, 1):
            p = self.doc.add_paragraph(f"{i}. {rec}")
            p.paragraph_format.left_indent = Inches(0.25)

        # Conclusion
        self._add_styled_heading('CONCLUSION', level=1, color=RGBColor(0, 0, 0))

        conclusion_text = f"""The evaluation framework successfully identified critical weaknesses in the human simulations:

• Significant underperformance vs. academic benchmark ({agg['semantic']['mean_cosine_similarity']:.0%} vs. 85% accuracy)
• Length mismatch is the most obvious tell ({agg['stylistic']['mean_length_ratio']:.1f}x too long)
• Lack of naturalness makes responses detectably artificial
• Category and individual variations suggest simulation quality is inconsistent

BOTTOM LINE: The simulations capture semantic content moderately well but fail to convincingly mimic human behavioral patterns, style, and naturalness. Substantial improvements are needed before deployment in production market research.

With targeted improvements to length calibration, naturalness injection, and personality tuning, the simulations can potentially close the {((0.85 - agg['semantic']['mean_cosine_similarity']) / 0.85 * 100):.0f}% performance gap and achieve human-level authenticity."""

        p = self.doc.add_paragraph(conclusion_text)
        p.paragraph_format.line_spacing = 1.3

        # Footer
        self.doc.add_paragraph()
        footer_p = self.doc.add_paragraph('_______________')
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer_text = self.doc.add_paragraph(f'Technical Report | Generated: {datetime.now().strftime("%B %d, %Y")} | Confidential')
        footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_text.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Save document
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_file))

        print(f"✅ Word report generated: {output_path}")
        return output_path


if __name__ == "__main__":
    generator = WordReportGenerator()
    generator.generate()
    print("\n📝 Editable Word document ready for customization!")
